"""
Document parsing utilities.

Supports:
- PDF (PyMuPDF → pdfplumber fallback)
- DOCX (paragraphs + tables)

Returns a list of page dicts:
    {
        "text":      str,
        "page":      int,
        "source":    str,
        "file_type": str,
        "file_path": str,
    }

Note on validation:
    Size and extension validation are the responsibility of the ingestion
    pipeline (validators.prepare_upload). This module only verifies the
    file exists before attempting to open it — a lightweight guard against
    race conditions between upload and parse.
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Any, Dict, List

logger = logging.getLogger(__name__)


# ──────────────────────────────────────────────────────────────
# Common helpers
# ──────────────────────────────────────────────────────────────

def _assert_exists(file_path: Path) -> None:
    """Raise FileNotFoundError if the file has disappeared since validation."""
    if not file_path.exists() or not file_path.is_file():
        raise FileNotFoundError(
            f"File not found or no longer accessible: {file_path}"
        )


def _make_page(
    *,
    text: str,
    page: int,
    file_path: Path,
) -> Dict[str, Any]:
    return {
        "text":      text.strip(),
        "page":      page,
        "source":    file_path.name,
        "file_type": file_path.suffix.lower().lstrip("."),
        "file_path": str(file_path.resolve()),
    }


# ──────────────────────────────────────────────────────────────
# PDF parsing
# ──────────────────────────────────────────────────────────────

def parse_pdf(file_path: "str | Path") -> List[Dict[str, Any]]:
    """
    Parse a PDF file and return a list of page dictionaries.

    Tries PyMuPDF first, then falls back to pdfplumber if PyMuPDF fails.

    Args:
        file_path: Path to the PDF file to parse.

    Returns:
        List of page dicts with keys: text, page, source, file_type, file_path.

    Raises:
        FileNotFoundError: If the file does not exist.
        RuntimeError: If both PDF parsers fail.
    """
    file_path = Path(file_path)
    _assert_exists(file_path)

    pages: List[Dict[str, Any]] = []

    # ── Try PyMuPDF first ─────────────────────────────────────
    try:
        import fitz  # PyMuPDF

        with fitz.open(file_path) as doc:
            for idx, page in enumerate(doc, start=1):
                text = page.get_text("text")
                if text and text.strip():
                    pages.append(_make_page(text=text, page=idx, file_path=file_path))

        if pages:
            logger.info("Parsed %d page(s) from '%s' via PyMuPDF.", len(pages), file_path.name)
            return pages

        logger.warning("PyMuPDF returned empty text for '%s'; falling back to pdfplumber.", file_path.name)

    except Exception as exc:
        logger.warning("PyMuPDF failed for '%s' (%s); falling back to pdfplumber.", file_path.name, exc)

    # ── Fallback: pdfplumber ──────────────────────────────────
    try:
        import pdfplumber

        with pdfplumber.open(file_path) as pdf:
            for idx, page in enumerate(pdf.pages, start=1):
                text = page.extract_text() or ""
                if text.strip():
                    pages.append(_make_page(text=text, page=idx, file_path=file_path))

        logger.info("Parsed %d page(s) from '%s' via pdfplumber.", len(pages), file_path.name)
        return pages

    except Exception as exc:
        logger.error("Both PDF parsers failed for '%s': %s", file_path.name, exc)
        raise RuntimeError(f"Cannot parse PDF '{file_path.name}'") from exc


# ──────────────────────────────────────────────────────────────
# DOCX parsing
# ──────────────────────────────────────────────────────────────

def parse_docx(file_path: "str | Path") -> List[Dict[str, Any]]:
    """
    Parse a DOCX file and return a list of synthetic page dicts.

    Paragraphs and table cells are extracted then grouped into synthetic
    pages of 30 blocks each for consistent downstream chunking.

    Args:
        file_path: Path to the DOCX file to parse.

    Returns:
        List of page dicts.

    Raises:
        FileNotFoundError: If the file does not exist.
    """
    from docx import Document

    file_path = Path(file_path)
    _assert_exists(file_path)

    doc = Document(file_path)
    blocks: List[str] = []

    # ── Paragraphs ────────────────────────────────────────────
    for p in doc.paragraphs:
        if p.text.strip():
            blocks.append(p.text.strip())

    # ── Tables ────────────────────────────────────────────────
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                blocks.append(row_text)

    # ── Group into synthetic pages ────────────────────────────
    GROUP_SIZE = 30
    pages: List[Dict[str, Any]] = []

    for idx in range(0, len(blocks), GROUP_SIZE):
        chunk = "\n".join(blocks[idx : idx + GROUP_SIZE])
        pages.append(_make_page(text=chunk, page=(idx // GROUP_SIZE) + 1, file_path=file_path))

    logger.info("Parsed %d page-group(s) from '%s'.", len(pages), file_path.name)
    return pages


# ──────────────────────────────────────────────────────────────
# Auto-dispatch
# ──────────────────────────────────────────────────────────────

def parse_document(file_path: "str | Path") -> List[Dict[str, Any]]:
    """Dispatch to the appropriate parser based on file extension."""
    file_path = Path(file_path)
    ext = file_path.suffix.lower()

    if ext == ".pdf":
        return parse_pdf(file_path)

    if ext in {".docx", ".doc"}:
        return parse_docx(file_path)

    raise ValueError(f"Unsupported file type '{ext}'. Supported: PDF, DOCX")
